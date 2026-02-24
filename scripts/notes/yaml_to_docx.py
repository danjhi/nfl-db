#!/usr/bin/env python3
"""Convert player_writeups.yaml to a Word docx for editing.

Format: Each player is a heading (Name - POS, TEAM) followed by the writeup paragraph.
"""

import sys, os, re
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'ids'))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

YAML_PATH = os.path.join(os.path.dirname(__file__), '..', '..', 'data', 'writeups', 'player_writeups.yaml')
DOCX_PATH = os.path.join(os.path.dirname(__file__), '..', '..', 'data', 'writeups', 'player_writeups.docx')


def parse_yaml(path):
    """Simple parser matching push_writeups.py pattern."""
    entries = []
    current = {}
    with open(path) as f:
        for line in f:
            line = line.rstrip('\n')
            if line.strip().startswith('- player_id:'):
                if current:
                    entries.append(current)
                current = {'player_id': line.split('"')[1]}
            elif line.strip().startswith('name:'):
                current['name'] = line.split('"')[1]
            elif line.strip().startswith('position:'):
                current['position'] = line.strip().split('position:')[1].strip()
            elif line.strip().startswith('team:'):
                current['team'] = line.strip().split('team:')[1].strip()
            elif line.strip().startswith('writeup:'):
                # Extract text between quotes
                m = re.search(r'writeup:\s*"(.*)"', line)
                if m:
                    current['writeup'] = m.group(1)
    if current:
        entries.append(current)
    return entries


def build_docx(entries, out_path):
    doc = Document()

    # Style tweaks
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Player Writeups', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'{len(entries)} players\n')

    for entry in entries:
        name = entry.get('name', '???')
        pos = entry.get('position', '?')
        team = entry.get('team', '?')
        writeup = entry.get('writeup', '').strip()

        # Heading: "Player Name - POS, TEAM"
        heading = doc.add_heading(f'{name} - {pos}, {team}', level=2)

        # Writeup paragraph
        if writeup:
            doc.add_paragraph(writeup)
        else:
            doc.add_paragraph('[No writeup]')

    doc.save(out_path)
    print(f'Saved {out_path} ({len(entries)} players)')


if __name__ == '__main__':
    entries = parse_yaml(YAML_PATH)
    build_docx(entries, DOCX_PATH)
